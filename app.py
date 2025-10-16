import base64
import binascii
import io
import os
from datetime import datetime
from typing import Optional

from flask import Flask, render_template, request, jsonify, send_file
from dotenv import load_dotenv
import requests
from docx import Document
from docx.shared import Inches

# .env 파일에서 환경 변수를 로드합니다.
load_dotenv()

app = Flask(__name__, template_folder='templates')

# OpenAI API 키를 환경 변수에서 가져옵니다.
API_KEY = os.getenv("OPENAI_API_KEY")
OPENAI_API_URL = "https://api.openai.com/v1/responses"


def _normalize_content_parts(content_parts):
    """Convert legacy chat-completions content parts into Responses API format."""
    normalized = []
    if not isinstance(content_parts, list):
        return normalized

    for part in content_parts:
        if not isinstance(part, dict):
            continue

        part_type = part.get("type")

        if part_type in {"text", "input_text"}:
            text_value = part.get("text")
            if isinstance(text_value, str):
                normalized.append({"type": "input_text", "text": text_value})
            continue

        if part_type in {"image_url", "input_image"}:
            image_url = part.get("image_url")
            if isinstance(image_url, dict):
                url_value = image_url.get("url")
            else:
                url_value = image_url
            if isinstance(url_value, str) and url_value:
                normalized.append({"type": "input_image", "image_url": url_value})
            continue

        # Passthrough any already-normalized content blocks.
        if part_type in {"input_audio", "input_video", "input_file"}:
            normalized.append(part)

    if not normalized:
        for part in content_parts:
            if isinstance(part, dict):
                normalized.append(part)

    return normalized


def _extract_text_from_response(response_payload):
    """Pull a best-effort assistant text string from a Responses API payload."""
    if not isinstance(response_payload, dict):
        return ""

    output_text = response_payload.get("output_text")
    if isinstance(output_text, str) and output_text.strip():
        return output_text

    collected = []
    output_items = response_payload.get("output")
    if isinstance(output_items, list):
        for item in output_items:
            if not isinstance(item, dict):
                continue
            if item.get("type") != "message":
                continue
            contents = item.get("content") or []
            if not isinstance(contents, list):
                continue
            for content in contents:
                if not isinstance(content, dict):
                    continue
                text_value = content.get("text")
                if isinstance(text_value, str) and text_value:
                    collected.append(text_value)

    if collected:
        return "\n".join(collected)

    # Fallback: if the upstream response still uses chat-completions layout, read it.
    choices = response_payload.get("choices")
    if isinstance(choices, list) and choices:
        first_choice = choices[0]
        if isinstance(first_choice, dict):
            message = first_choice.get("message")
            if isinstance(message, dict):
                content = message.get("content")
                if isinstance(content, str):
                    return content

    return ""

@app.route('/')
def index():
    """메인 웹페이지를 렌더링합니다."""
    return render_template('index.html')

@app.route('/analyze', methods=['POST'])
def analyze():
    """
    프론트엔드로부터 요청을 받아 OpenAI API를 호출하고 결과를 반환합니다.
    """
    if not API_KEY:
        print("🛑 오류: .env 파일에 OPENAI_API_KEY가 설정되지 않았습니다.")
        return jsonify({"error": ".env 파일에 OPENAI_API_KEY가 설정되지 않았습니다."}), 500

    data = request.json
    
    content_parts = data.get("content_parts", [])
    normalized_content = _normalize_content_parts(content_parts)

    # 프론트엔드에서 받은 데이터를 기반으로 OpenAI에 보낼 payload를 구성합니다.
    requested_model = data.get("model", "gpt-5-mini")
    model_presets = {
        "gpt-5-mini": {
            "max_output_tokens": 4096,
            "temperature": 0.2,
            "top_p": 0.8,
        },
        "gpt-4o": {
            "max_output_tokens": 3072,
            "temperature": 0.6,
            "top_p": 0.9,
        },
    }

    if requested_model not in model_presets:
        print(f"⚠️ 지원하지 않는 모델이 요청되었습니다: {requested_model}. gpt-5-mini로 대체합니다.")

    model = requested_model if requested_model in model_presets else "gpt-5-mini"
    model_payload_options = model_presets[model]

    payload = {
        # gpt-5-mini는 텍스트와 이미지 URL이 혼합된 메시지를 처리할 수 있는 멀티모달 모델입니다.
        "model": model,
        "input": [{
            "role": "user",
            "content": normalized_content,
        }],
        **model_payload_options,
    }

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {API_KEY}",
    }

    try:
        print(
            "🚀 OpenAI API에 분석을 요청합니다... "
            f"(model={model}, temperature={model_payload_options['temperature']}, "
            f"max_output_tokens={model_payload_options['max_output_tokens']}, top_p={model_payload_options['top_p']})"
        )
        response = requests.post(OPENAI_API_URL, headers=headers, json=payload)
        response.raise_for_status()  # HTTP 오류가 발생하면 예외를 발생시킵니다.
        print("✅ OpenAI API로부터 응답을 받았습니다.")
        response_payload = response.json()
        analysis_text = _extract_text_from_response(response_payload)
        return jsonify({
            "choices": [{"message": {"content": analysis_text}}],
            "raw_response": response_payload,
        })
        
    except requests.exceptions.RequestException as e:
        # API 호출 실패 시 에러 메시지를 JSON 형식으로 반환합니다.
        error_message = str(e)
        if e.response:
            try:
                error_detail = e.response.json()
                error_message = error_detail.get("error", {}).get("message", "알 수 없는 오류")
            except ValueError:
                error_message = e.response.text
        
        print(f"🛑 OpenAI API 호출 실패: {error_message}")
        return jsonify({"error": f"OpenAI API 호출 실패: {error_message}"}), 500


def _decode_data_url(data_url: Optional[str]) -> Optional[bytes]:
    """Decode a base64 data URL into raw bytes."""
    if not data_url or not data_url.startswith("data:"):
        return None
    try:
        header, encoded = data_url.split(",", 1)
    except ValueError:
        return None
    if ";base64" not in header:
        return None
    try:
        return base64.b64decode(encoded)
    except (binascii.Error, ValueError):
        return None


def _sanitize_filename(filename: str) -> str:
    base_name = filename.rsplit('.', 1)[0]
    safe = ''.join(ch for ch in base_name if ch.isalnum() or ch in (' ', '_', '-')).strip()
    return safe or "analysis_report"


@app.route('/create-report', methods=['POST'])
def create_report():
    if not request.is_json:
        return jsonify({"error": "유효하지 않은 요청 형식입니다. JSON 데이터를 전송해주세요."}), 400

    payload = request.get_json(silent=True) or {}
    title = payload.get("title", "보고서")
    global_summary = payload.get("global_summary", "")
    analysis_results = payload.get("analysis_results", [])
    if not isinstance(analysis_results, list):
        analysis_results = []

    document = Document()

    document.add_heading(title.replace('.pdf', ''), level=0)
    document.add_paragraph(f"생성일: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    document.add_page_break()
    document.add_heading("Executive Summary", level=1)
    for line in str(global_summary).split('\n'):
        document.add_paragraph(line)

    document.add_page_break()
    document.add_heading("상세 분석 (Detailed Analysis)", level=1)

    for index, result in enumerate(analysis_results, start=1):
        group = result.get("group") if isinstance(result, dict) else {}
        if not isinstance(group, dict):
            group = {}
        group_id = group.get("id", index)
        pages = group.get("pages") or []
        pages_text = ", ".join(str(page) for page in pages) if pages else "N/A"
        document.add_heading(f"그룹 {group_id} (페이지: {pages_text})", level=2)

        intent = group.get("intent", "")
        intent_paragraph = document.add_paragraph()
        intent_run = intent_paragraph.add_run("분석 초점: ")
        intent_run.bold = True
        intent_paragraph.add_run(str(intent))

        images = result.get("images") if isinstance(result, dict) else []
        for image_data_url in images or []:
            image_bytes = _decode_data_url(image_data_url)
            if not image_bytes:
                continue
            image_stream = io.BytesIO(image_bytes)
            try:
                document.add_picture(image_stream, width=Inches(6))
            except Exception as picture_error:  # pylint: disable=broad-except
                document.add_paragraph(f"[이미지 추가 실패: {picture_error}]")

        analysis_text = result.get("analysis", "") if isinstance(result, dict) else ""
        for line in str(analysis_text).split('\n'):
            document.add_paragraph(line)

        document.add_paragraph()

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    filename = f"{_sanitize_filename(str(title))}_analysis_report.docx"

    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

if __name__ == '__main__':
    port = 5001
    print(f"✅ 서버를 시작합니다. http://127.0.0.1:{port} 에서 접속하세요.")
    app.run(debug=True, port=port)
